from docx import Document
from docx.shared import Cm

document = Document()

# profile picture
document.add_picture('me.png', width = Cm(10))

# name, phone number and email
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name + " " + phone_number + " " + email)

# about me 
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself? '))

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences?')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

# list of skills
document.add_heading('Skills')

has_skills = input('What are your skills? ')
document.add_paragraph(has_skills, style = 'List Bullet')

while True:
    has_more_skills = input('Do you have more skills?')
    if has_more_skills.lower() == 'yes':
        has_more_skills = input('What are your other skills? ')
        document.add_paragraph(has_more_skills, style = 'List Bullet')
    else:
        break

# save document
document.save('Cv.docx')